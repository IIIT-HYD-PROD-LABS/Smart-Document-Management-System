"""Build PRODUCT_FEATURES.docx from the markdown source.

No external binary deps (no pandoc). Walks markdown-it-py tokens and
emits python-docx paragraphs / lists / tables / code runs. Mirrors the
in-tree pattern used by build_tech_pdf.py (markdown-it-py + Chrome
headless for the PDF variant).

Run:  python scripts/build_features_docx.py
Output: docs/exports/PRODUCT_FEATURES.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from markdown_it import MarkdownIt

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "reference" / "PRODUCT_FEATURES.md"
DST = ROOT / "docs" / "exports" / "PRODUCT_FEATURES.docx"

# Heading sizes (pt) by markdown level. h1 is the doc title; h2-h4 are
# the bulk of the body. We cap at h6 for safety even though the doc
# only goes to h4.
HEADING_PT = {1: 22, 2: 18, 3: 14, 4: 12, 5: 11, 6: 11}

# Brand-blue used for headings — matches TaxSync's --accent #2563eb.
ACCENT = RGBColor(0x25, 0x63, 0xEB)


def add_inline(paragraph, children):
    """Walk inline tokens (token.children) and append runs to `paragraph`.

    Tracks open/close depth for **strong** and *em* so nested marks work.
    Inline `code` becomes a Consolas run. Links render their visible text;
    URLs are omitted in this v1 to keep the body readable when printed.
    """
    bold_depth = 0
    em_depth = 0
    for c in children or []:
        ttype = c.type
        if ttype == "text":
            run = paragraph.add_run(c.content)
            if bold_depth:
                run.bold = True
            if em_depth:
                run.italic = True
        elif ttype == "code_inline":
            run = paragraph.add_run(c.content)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif ttype == "strong_open":
            bold_depth += 1
        elif ttype == "strong_close":
            bold_depth -= 1
        elif ttype == "em_open":
            em_depth += 1
        elif ttype == "em_close":
            em_depth -= 1
        elif ttype in ("softbreak", "hardbreak"):
            paragraph.add_run(" ")
        # link_open / link_close: text content is emitted by adjacent
        # `text` children; we drop the URL to keep print output clean.


def emit_list_item(doc, tokens, start, list_style):
    """Emit one list_item as a single paragraph in `list_style`.

    Returns the index AFTER the matching list_item_close. Markdown-it
    represents nested paragraphs inside list items, so we collect the
    first paragraph's inline content (matches how the source doc is
    written — bullets are single-line)."""
    j = start + 1
    while j < len(tokens) and tokens[j].type != "list_item_close":
        if tokens[j].type == "paragraph_open" and tokens[j + 1].type == "inline":
            p = doc.add_paragraph(style=list_style)
            add_inline(p, tokens[j + 1].children)
            j += 3  # paragraph_open, inline, paragraph_close
        else:
            j += 1
    return j + 1  # past list_item_close


def emit_table(doc, tokens, start):
    """Read a Markdown-it table block into a Word table, returning the
    index AFTER the matching table_close."""
    rows = []  # list of (is_header_row, [cell_inline_children])
    j = start + 1
    in_header = False
    current_row = None
    while j < len(tokens) and tokens[j].type != "table_close":
        t = tokens[j]
        if t.type == "thead_open":
            in_header = True
        elif t.type == "thead_close":
            in_header = False
        elif t.type == "tr_open":
            current_row = []
        elif t.type == "tr_close":
            if current_row is not None:
                rows.append((in_header, current_row))
                current_row = None
        elif t.type in ("th_open", "td_open"):
            if tokens[j + 1].type == "inline":
                current_row.append(tokens[j + 1].children)
            else:
                current_row.append([])
            j += 2  # th_open + inline (close consumed below)
        j += 1

    if not rows:
        return j + 1

    ncols = max(len(cells) for _, cells in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    for ri, (is_header, cells) in enumerate(rows):
        row = table.rows[ri]
        for ci in range(ncols):
            cell = row.cells[ci]
            cell.text = ""  # python-docx adds an empty paragraph; reuse it
            p = cell.paragraphs[0]
            children = cells[ci] if ci < len(cells) else []
            add_inline(p, children)
            if is_header:
                for run in p.runs:
                    run.bold = True

    return j + 1  # past table_close


def emit_fence(doc, content):
    """Render a code block as a single mono-spaced paragraph (joining the
    lines with line-breaks via add_break) so the visual block stays
    together rather than fragmenting across N paragraphs."""
    p = doc.add_paragraph()
    lines = content.rstrip("\n").split("\n") or [""]
    for idx, line in enumerate(lines):
        run = p.add_run(line or " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        if idx < len(lines) - 1:
            run.add_break()


def main() -> None:
    text = SRC.read_text(encoding="utf-8")
    parser = MarkdownIt("commonmark", {"html": False, "linkify": False}).enable(
        "table",
    )
    tokens = parser.parse(text)

    doc = Document()
    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(11)

    i = 0
    while i < len(tokens):
        t = tokens[i]
        ttype = t.type

        if ttype == "heading_open":
            level = int(t.tag[1])  # 'h2' -> 2
            heading = doc.add_heading(level=min(level, 9))
            for run in heading.runs:
                run.text = ""
            inline = tokens[i + 1]
            add_inline(heading, inline.children)
            for run in heading.runs:
                run.font.size = Pt(HEADING_PT.get(level, 11))
                run.font.color.rgb = ACCENT
                if level <= 2:
                    run.bold = True
            i += 3

        elif ttype == "paragraph_open":
            inline = tokens[i + 1]
            p = doc.add_paragraph()
            add_inline(p, inline.children)
            i += 3

        elif ttype == "bullet_list_open":
            i += 1
            while i < len(tokens) and tokens[i].type != "bullet_list_close":
                if tokens[i].type == "list_item_open":
                    i = emit_list_item(doc, tokens, i, list_style="List Bullet")
                else:
                    i += 1
            i += 1

        elif ttype == "ordered_list_open":
            i += 1
            while i < len(tokens) and tokens[i].type != "ordered_list_close":
                if tokens[i].type == "list_item_open":
                    i = emit_list_item(doc, tokens, i, list_style="List Number")
                else:
                    i += 1
            i += 1

        elif ttype == "table_open":
            i = emit_table(doc, tokens, i)

        elif ttype == "fence":
            emit_fence(doc, t.content)
            i += 1

        elif ttype == "hr":
            # Visual separator — a thin paragraph with em-dash run
            sep = doc.add_paragraph()
            run = sep.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1

        else:
            i += 1

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DST)
    print(f"Wrote {DST.relative_to(ROOT)} ({DST.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
