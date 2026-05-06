"""Convert STATUS_REPORT.md → Smart_Document_Management_System_Status_Report.docx.

Lightweight Markdown → Word converter via python-docx. Handles the subset
of Markdown actually used in our status reports:

  - ATX headings (# .. ######)
  - Horizontal rules (---)
  - Unordered list items (- foo)
  - Pipe tables
  - Code spans (`foo`) → monospace inline runs
  - Bold (**foo**) and italic (*foo*) → bold/italic runs
  - Plain paragraphs

This isn't a general-purpose converter — it's tuned for the structure of
our own status doc. We control the input, so we don't need to handle
weird edge cases.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor


# ───────────────────────────────────────────────────────────────────
# Inline parser — handles **bold**, *italic*, and `code` runs
# ───────────────────────────────────────────────────────────────────
INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)"
)


def _add_inline_runs(paragraph, text: str) -> None:
    """Split `text` into bold/italic/code/plain runs and append to paragraph."""
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            paragraph.add_run(part)


# ───────────────────────────────────────────────────────────────────
# Block-level parser
# ───────────────────────────────────────────────────────────────────
def _add_horizontal_rule(paragraph) -> None:
    """Insert a hairline horizontal rule below the paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B0B0B0")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _heading_style(level: int) -> str:
    return {
        1: "Title",
        2: "Heading 1",
        3: "Heading 2",
        4: "Heading 3",
        5: "Heading 4",
        6: "Heading 5",
    }.get(level, "Heading 5")


def _flush_table(doc, rows: list[list[str]]) -> None:
    """Emit a Word table from accumulated pipe-table rows."""
    if not rows:
        return
    # First row is header; second row (separator) was already discarded.
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.rows[ri].cells[ci]
            # Clear default paragraph and add styled content
            cell.paragraphs[0].text = ""
            p = cell.paragraphs[0]
            _add_inline_runs(p, cell_text.strip())
            if ri == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()  # spacer after table


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    md = md_path.read_text(encoding="utf-8")
    doc = Document()

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md.splitlines()
    i = 0
    pending_table: list[list[str]] = []

    while i < len(lines):
        line = lines[i]

        # Pipe-table detection: `| a | b |` followed by separator line
        if line.strip().startswith("|") and i + 1 < len(lines) and re.match(
            r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", lines[i + 1]
        ):
            pending_table = []
            # Header row
            cells = [c.strip() for c in line.strip("|").split("|")]
            pending_table.append(cells)
            i += 2  # skip header + separator
            # Body rows
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip("|").split("|")]
                pending_table.append(cells)
                i += 1
            _flush_table(doc, pending_table)
            pending_table = []
            continue

        # Horizontal rule
        if re.match(r"^\s*-{3,}\s*$", line):
            p = doc.add_paragraph()
            _add_horizontal_rule(p)
            i += 1
            continue

        # ATX heading
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            heading = doc.add_heading(level=level if level <= 4 else 4)
            heading.style = doc.styles[_heading_style(level)]
            run = heading.add_run(text)
            if level == 1:
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor(0x10, 0x10, 0x10)
            elif level == 2:
                run.font.size = Pt(15)
                run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
            else:
                run.font.size = Pt(13)
            i += 1
            continue

        # Unordered list item
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            text = m.group(2).strip()
            level = min(indent // 2, 4)
            try:
                p = doc.add_paragraph(style="List Bullet" + (f" {level + 1}" if level > 0 else ""))
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, text)
            i += 1
            continue

        # Blank line
        if line.strip() == "":
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        _add_inline_runs(p, line)
        i += 1

    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


def main() -> int:
    project_root = Path(__file__).resolve().parent.parent
    md = project_root / "STATUS_REPORT.md"
    if not md.exists():
        print(f"FATAL: {md} not found", file=sys.stderr)
        return 1

    # Two outputs: the canonical clean filename + the legacy
    # `.md (1) (1).docx` filename the user already has open.
    canonical = project_root / "Smart_Document_Management_System_Status_Report.docx"
    legacy = project_root / "Smart_Document_Management_System_Status_Report.md (1) (1).docx"

    md_to_docx(md, canonical)
    # Mirror to the legacy filename so the user's existing reference doesn't go stale
    md_to_docx(md, legacy)
    return 0


if __name__ == "__main__":
    sys.exit(main())
