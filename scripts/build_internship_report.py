"""Build a single-page internship report DOCX from inline content.

Writes /home/sraav/Desktop/Smart_Docs_Prod_Labs/Internship_Report_TaxSync.docx

Tone: professional, plain language a non-technical mentor can follow.
Layout: 1 page, A4, generous margins for readability.

Style rule: NO em-dashes or en-dashes anywhere. The user explicitly flagged
those as an AI tell. Use commas, periods, parentheses, or "to" instead.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm


OUT = Path("/home/sraav/Desktop/Smart_Docs_Prod_Labs/Internship_Report_TaxSync.docx")

ACCENT = RGBColor(0x1F, 0x3D, 0x7A)
SUBTLE = RGBColor(0x55, 0x5B, 0x66)
BODY = RGBColor(0x1A, 0x1F, 0x29)


def _set_font(run, size_pt, *, bold=False, color=BODY, italic=False, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    _set_font(r, 16, bold=True, color=ACCENT)


def add_meta(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(1)
    r = p.add_run(text)
    _set_font(r, 9.5, color=SUBTLE, italic=True)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    _set_font(r, 11, bold=True, color=ACCENT)


def add_para(doc, text, *, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    _set_font(r, size, color=BODY)


def add_bullet(doc, lead, rest, *, size=10.5):
    """Bullet with a short bold lead-in followed by mentor-friendly prose."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    if lead:
        lead_run = p.add_run(lead)
        _set_font(lead_run, size, bold=True, color=BODY)
    rest_run = p.add_run(rest)
    _set_font(rest_run, size, color=BODY)


def build() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # ── Header block ─────────────────────────────────────────────────
    add_title(doc, "Internship Report: TaxSync (Smart Document Management System)")
    add_meta(
        doc,
        "Name: P CH NVS Sravan Kumar   |   "
        "Organisation: Product Labs, IIIT Hyderabad   |   "
        "Period: January to April 2026",
    )

    # ── Project overview ────────────────────────────────────────────
    add_heading(doc, "Project Overview")
    add_para(
        doc,
        "TaxSync is a web platform built for Indian tax and compliance "
        "teams. Users upload any document such as PDFs, scanned images, "
        "or Word files, and the system extracts the text, classifies it "
        "using a machine learning model, and makes it searchable. On top "
        "of that, the platform tracks regulatory notices from GST, "
        "Income Tax, MCA, RBI, and SEBI through a four stage approval "
        "workflow with automatic deadline alerts and per client "
        "compliance reports.",
    )

    # ── Contributions ───────────────────────────────────────────────
    add_heading(doc, "My Contributions")
    add_bullet(
        doc,
        "Full stack development. ",
        "Designed and shipped features across the backend, the web app, "
        "and the machine learning pipeline that processes uploaded documents.",
    )
    add_bullet(
        doc,
        "Document workflow. ",
        "Built the upload, text extraction (OCR), classification, and "
        "search pipeline. Added document sharing with permission roles, "
        "version history with rollback, and an audit trail.",
    )
    add_bullet(
        doc,
        "Compliance module. ",
        "Implemented the multi tenant compliance tracker, with each "
        "client's data isolated at the database level. It includes a "
        "four stage approval workflow, automatic deadline alerts at 7, "
        "3, and 1 days before due, and per client reports.",
    )
    add_bullet(
        doc,
        "AI assistant. ",
        "Integrated a bring your own key AI assistant supporting "
        "Anthropic Claude and Google Gemini, scope limited to regulatory "
        "and finance work so the system stays on task.",
    )
    add_bullet(
        doc,
        "Performance and security. ",
        "Tuned the database access pattern to roughly halve API response "
        "time. Added security hardening: encrypted credentials, login "
        "with Google or Microsoft accounts, role based access, and an "
        "immutable audit log.",
    )

    # ── Tech stack ──────────────────────────────────────────────────
    add_heading(doc, "Technology Stack")
    add_para(
        doc,
        "Python, FastAPI, SQLAlchemy, PostgreSQL (hosted on Supabase), "
        "Redis, Celery, scikit learn, Tesseract OCR, Next.js, React, "
        "TypeScript, Tailwind CSS, Docker, and GitHub Actions for "
        "continuous integration.",
    )

    # ── Outcomes ────────────────────────────────────────────────────
    add_heading(doc, "Key Outcomes")
    add_bullet(
        doc,
        "",
        "Production shipped v1.0 through v2.1.1 across approximately 16 development phases.",
    )
    add_bullet(
        doc,
        "",
        "More than 502 backend automated tests kept passing on every commit.",
    )
    add_bullet(
        doc,
        "",
        "Roughly 2x faster API health check and 3x faster multi query "
        "endpoints after the database access tuning.",
    )
    add_bullet(
        doc,
        "",
        "Navigation consolidated from 19 items to 14 across five groups, "
        "based on real usage observed during demos.",
    )

    # ── Skills ──────────────────────────────────────────────────────
    add_heading(doc, "Skills Developed")
    add_para(
        doc,
        "End to end SaaS engineering across the full stack, multi tenant "
        "database design, OCR and machine learning workflows, secure "
        "login and key handling, performance profiling, code review with "
        "AI agents, and the discipline of shipping production features "
        "while keeping a large test suite green.",
    )

    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
