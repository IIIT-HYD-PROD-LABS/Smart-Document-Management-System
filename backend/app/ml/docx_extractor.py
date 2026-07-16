"""DOCX text extraction module."""
import io
import zipfile
import structlog
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.ml.errors import ExtractionEngineError

MAX_TEXT_LENGTH = 500_000
# S2: python-docx unzips the .docx container with no decompressed-size limit,
# so a small zip bomb can expand to gigabytes. Reject any container whose
# uncompressed payload (per-member or total) exceeds this cap before parsing.
_MAX_DOCX_UNCOMPRESSED_BYTES = 200 * 1024 * 1024

logger = structlog.stdlib.get_logger()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file, including paragraphs and tables."""
    try:
        # S2: cap the decompressed payload before python-docx unzips the
        # container, so a zip bomb can't expand to gigabytes and exhaust memory.
        # A non-zip container raises BadZipFile here and is handled as a corrupt
        # docx by the except clause below.
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            total_uncompressed = 0
            for info in zf.infolist():
                total_uncompressed += info.file_size
                if (
                    info.file_size > _MAX_DOCX_UNCOMPRESSED_BYTES
                    or total_uncompressed > _MAX_DOCX_UNCOMPRESSED_BYTES
                ):
                    raise ExtractionEngineError("docx_too_large", retryable=False)

        doc = DocxDocument(io.BytesIO(file_bytes))
        parts = []
        total_len = 0
        truncated = False
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                total_len += len(text)
                if total_len > MAX_TEXT_LENGTH:
                    truncated = True
                    break
                parts.append(text)
        if not truncated:
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        row_text = " | ".join(cells)
                        total_len += len(row_text)
                        if total_len > MAX_TEXT_LENGTH:
                            truncated = True
                            break
                        parts.append(row_text)
                if truncated:
                    break
        result = "\n".join(parts)
        if truncated:
            logger.warning("docx_text_truncated", chars=len(result), limit=MAX_TEXT_LENGTH)
        else:
            logger.info("docx_text_extracted", chars=len(result))
        return result
    except ExtractionEngineError:
        # S2: our own size-cap rejection must propagate; the broad Exception
        # handler below would otherwise swallow it into an empty result.
        raise
    except (zipfile.BadZipFile, PackageNotFoundError) as e:
        # Corrupt / not-a-real-docx container. Surfacing "" would mark the
        # document COMPLETED with no content; raise so the task layer marks
        # it FAILED. Non-retryable: a corrupt file never decodes on retry.
        logger.error("docx_extraction_failed_bad_zip", error=str(e))
        raise ExtractionEngineError("docx_corrupt", retryable=False) from e
    except ValueError as e:
        logger.error("docx_extraction_failed_value_error", error=str(e))
        return ""
    except Exception as e:
        logger.error("docx_extraction_failed", error=str(e))
        return ""
