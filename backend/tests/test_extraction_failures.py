"""Regression tests for silent-failure bugs in the extraction pipeline.

Covers the audit findings C1/H1/H2/H3/M4 and the DOCX corrupt-file gap.

These are pure unit tests: heavy engines (tesseract, the sklearn model,
pdfplumber) are mocked. No DB or Celery broker is required.
"""

import io
from unittest.mock import MagicMock, patch, PropertyMock

import pytest
from celery.app.task import Task as _CeleryBaseTask
from PIL import Image
from pytesseract import TesseractNotFoundError

from app.ml.errors import ExtractionEngineError
from app.models.document import DocumentStatus, DocumentCategory
from app.tasks.document_tasks import process_document_task


# ---------------------------------------------------------------------------
# Shared task-harness helpers (mirrors tests/test_document_tasks.py)
# ---------------------------------------------------------------------------

_TASK = "app.tasks.document_tasks"
_STORAGE = "app.services.storage_service"


class _TaskSelfContext:
    def __init__(self, retries: int = 0, max_retries: int = 3):
        self.retries = retries
        self.max_retries = max_retries
        self._patches = []

    def __enter__(self):
        mock_request = MagicMock()
        mock_request.retries = self.retries
        p1 = patch.object(_CeleryBaseTask, "request", new_callable=PropertyMock, return_value=mock_request)
        p2 = patch.object(process_document_task, "update_state", MagicMock())
        p3 = patch.object(process_document_task, "max_retries", self.max_retries)
        p4 = patch.object(process_document_task, "retry", MagicMock(side_effect=Exception("retry called")))
        p5 = patch.object(process_document_task, "MaxRetriesExceededError", Exception)
        self._patches = [p1, p2, p3, p4, p5]
        for p in self._patches:
            p.start()
        return self

    def __exit__(self, *args):
        for p in reversed(self._patches):
            p.stop()


def _make_mock_document(**overrides):
    doc = MagicMock()
    doc.id = overrides.get("id", 1)
    doc.file_path = overrides.get("file_path", "/uploads/test.pdf")
    doc.file_type = overrides.get("file_type", "pdf")
    doc.notice_id = None
    doc.status = DocumentStatus.PENDING
    doc.extracted_text = None
    doc.category = DocumentCategory.UNKNOWN
    doc.confidence_score = 0.0
    doc.extracted_metadata = None
    doc.ai_summary = None
    doc.ai_extracted_fields = None
    doc.ai_provider = None
    doc.ai_extraction_status = None
    return doc


def _make_mock_db(doc):
    mock_db = MagicMock()
    mock_db.query.return_value.filter.return_value.first.return_value = doc
    return mock_db


def _run_task_with(doc, *, extract_side_effect=None, extract_return=None,
                   truncated=False, file_type="pdf"):
    """Run process_document_task with extract_and_classify mocked."""
    doc.file_type = file_type
    mock_db = _make_mock_db(doc)
    with patch(f"{_TASK}.SessionLocal", return_value=mock_db), \
         patch(f"{_STORAGE}._validate_path_inside_upload_dir", return_value="/uploads/test"), \
         patch(f"{_TASK}.os.path.exists", return_value=True), \
         patch(f"{_TASK}.os.path.getsize", return_value=1024), \
         patch(f"{_TASK}.settings") as mock_settings, \
         patch(f"{_TASK}.consume_scanned_pdf_truncated", return_value=truncated), \
         patch("builtins.open", MagicMock()), \
         patch(f"{_TASK}.extract_and_classify") as mock_extract, \
         _TaskSelfContext():
        mock_settings.MAX_FILE_SIZE_MB = 50
        if extract_side_effect is not None:
            mock_extract.side_effect = extract_side_effect
        else:
            mock_extract.return_value = extract_return
        result = process_document_task.run(1)
    return result


# ---------------------------------------------------------------------------
# C1 — empty extraction from a non-.txt source must NOT end COMPLETED
# ---------------------------------------------------------------------------

class TestEmptyExtractionNotCompleted:

    def test_empty_pdf_extraction_marks_failed_not_completed(self):
        doc = _make_mock_document()
        result = _run_task_with(doc, extract_return=("", "unknown", 0.0), file_type="pdf")
        assert result["status"] == "failed"
        assert result["error"] == "no_text"
        assert doc.status == DocumentStatus.FAILED
        assert doc.ai_extraction_status == "no_text"

    def test_whitespace_only_extraction_marks_failed(self):
        doc = _make_mock_document()
        result = _run_task_with(doc, extract_return=("   \n  ", "unknown", 0.0), file_type="png")
        assert result["status"] == "failed"
        assert doc.status == DocumentStatus.FAILED

    def test_empty_txt_body_still_completes(self):
        """A .txt source is legitimately allowed to be empty (Gmail body)."""
        doc = _make_mock_document()
        result = _run_task_with(doc, extract_return=("", "unknown", 0.0), file_type="txt")
        assert result["status"] == "completed"
        assert doc.status == DocumentStatus.COMPLETED


# ---------------------------------------------------------------------------
# C1/H1/H2 — ExtractionEngineError routing in the task layer
# ---------------------------------------------------------------------------

class TestExtractionEngineErrorRouting:

    def test_retryable_engine_error_enters_retry_path(self):
        """Retryable engine error (e.g. tesseract down) funnels into the
        retry machinery and never commits COMPLETED.

        Under the test harness ``self.retry`` is mocked to raise, which the
        patched ``MaxRetriesExceededError`` catches, so the terminal state is
        FAILED -- the key contract is that ``self.retry`` was reached (the
        engine error was treated as transient) and the doc never COMPLETED.
        """
        doc = _make_mock_document()
        err = ExtractionEngineError("tesseract_not_found", retryable=True)
        # retry() is mocked to raise Exception("retry called"); reaching it
        # proves the error funneled into the retry path, not COMPLETED.
        with pytest.raises(Exception, match="retry called"):
            _run_task_with(doc, extract_side_effect=err, file_type="png")
        assert doc.status != DocumentStatus.COMPLETED
        assert doc.status == DocumentStatus.FAILED

    def test_non_retryable_engine_error_marks_failed(self):
        """Non-retryable engine error (corrupt container) fails permanently."""
        doc = _make_mock_document()
        err = ExtractionEngineError("docx_corrupt", retryable=False)
        result = _run_task_with(doc, extract_side_effect=err, file_type="docx")
        assert result["status"] == "failed"
        assert result["error"] == "docx_corrupt"
        assert doc.status == DocumentStatus.FAILED
        assert doc.ai_extraction_status == "failed"


# ---------------------------------------------------------------------------
# M4 — scanned-PDF OCR cap surfaces incomplete flag
# ---------------------------------------------------------------------------

class TestScannedPdfTruncationVisible:

    def test_truncated_scanned_pdf_sets_incomplete_status(self):
        doc = _make_mock_document()
        result = _run_task_with(
            doc,
            extract_return=("page text " * 10, "tax", 0.8),
            truncated=True,
            file_type="pdf",
        )
        assert result["status"] == "completed"
        assert doc.status == DocumentStatus.COMPLETED
        assert doc.ai_extraction_status == "incomplete_scanned_pdf"

    def test_non_truncated_pdf_keeps_normal_status(self):
        doc = _make_mock_document()
        result = _run_task_with(
            doc,
            extract_return=("page text " * 10, "tax", 0.8),
            truncated=False,
            file_type="pdf",
        )
        assert result["status"] == "completed"
        assert doc.ai_extraction_status != "incomplete_scanned_pdf"


# ---------------------------------------------------------------------------
# H1 — missing tesseract -> ExtractionEngineError (retryable)
# ---------------------------------------------------------------------------

class TestTesseractMissingRaises:

    def test_extract_text_from_image_raises_engine_error(self):
        from app.ml import ocr
        with patch.object(ocr.cv2, "imdecode", return_value=MagicMock(shape=(10, 10, 3))), \
             patch.object(ocr, "preprocess_image", return_value=MagicMock()), \
             patch.object(ocr.pytesseract, "image_to_string",
                          side_effect=TesseractNotFoundError()):
            with pytest.raises(ExtractionEngineError) as exc:
                ocr.extract_text_from_image(b"\x89PNG fake bytes")
        assert exc.value.retryable is True
        assert "tesseract" in exc.value.reason

    def test_blank_image_still_returns_empty_string(self):
        """A genuinely text-free image returns '' (no raise)."""
        from app.ml import ocr
        with patch.object(ocr.cv2, "imdecode", return_value=MagicMock(shape=(10, 10, 3))), \
             patch.object(ocr, "preprocess_image", return_value=MagicMock()), \
             patch.object(ocr.pytesseract, "image_to_string", return_value="   "):
            result = ocr.extract_text_from_image(b"fake")
        assert result == ""


# ---------------------------------------------------------------------------
# H2 — missing model artifacts -> ExtractionEngineError
# ---------------------------------------------------------------------------

class TestMissingModelArtifacts:

    def test_classify_document_raises_when_artifacts_missing(self):
        from app.ml import classifier
        # Reset the lazy globals so _load_model re-evaluates file existence.
        classifier._model = None
        classifier._vectorizer = None
        text = "x" * 100  # long enough to pass the length gate
        with patch.object(classifier.os.path, "exists", return_value=False):
            with pytest.raises(ExtractionEngineError) as exc:
                classifier.classify_document(text)
        assert exc.value.reason == "model_artifacts_missing"

    def test_short_text_returns_unknown_without_loading_model(self):
        """Genuine low-signal input still returns unknown, never raises."""
        from app.ml import classifier
        classifier._model = None
        classifier._vectorizer = None
        cat, conf = classifier.classify_document("tiny")
        assert cat == "unknown"
        assert conf == 0.0


# ---------------------------------------------------------------------------
# H3 — multi-page TIFF extracts text from ALL frames
# ---------------------------------------------------------------------------

def _make_two_frame_tiff() -> bytes:
    """Synthesize a real 2-frame TIFF in memory."""
    frame1 = Image.new("RGB", (40, 40), color="white")
    frame2 = Image.new("RGB", (40, 40), color="white")
    buf = io.BytesIO()
    frame1.save(buf, format="TIFF", save_all=True, append_images=[frame2])
    return buf.getvalue()


class TestMultiPageTiff:

    def test_all_frames_are_ocrd(self):
        from app.ml import ocr
        tiff_bytes = _make_two_frame_tiff()
        # OCR each frame deterministically so we can assert all frames ran.
        calls = {"n": 0}

        def _fake_pil_ocr(pil_image):
            calls["n"] += 1
            return f"frame-{calls['n']}-text"

        with patch.object(ocr, "extract_text_from_pil_image", side_effect=_fake_pil_ocr):
            text = ocr.extract_text_from_tiff(tiff_bytes)

        assert calls["n"] == 2, "both TIFF frames must be OCR'd"
        assert "frame-1-text" in text
        assert "frame-2-text" in text

    def test_extract_and_classify_routes_tiff_to_frame_iterator(self):
        from app.ml import classifier
        with patch("app.ml.ocr.extract_text_from_tiff", return_value="multi page text") as m_tiff, \
             patch("app.ml.ocr.extract_text_from_image", return_value="single") as m_img, \
             patch.object(classifier, "classify_document", return_value=("tax", 0.9)):
            text, cat, conf = classifier.extract_and_classify(b"fake-tiff", "tiff")
        m_tiff.assert_called_once()
        m_img.assert_not_called()
        assert text == "multi page text"

    def test_tiff_engine_error_propagates(self):
        """If tesseract is missing, the per-frame error propagates out."""
        from app.ml import ocr
        tiff_bytes = _make_two_frame_tiff()
        with patch.object(ocr, "extract_text_from_pil_image",
                          side_effect=ExtractionEngineError("tesseract_not_found")):
            with pytest.raises(ExtractionEngineError):
                ocr.extract_text_from_tiff(tiff_bytes)


# ---------------------------------------------------------------------------
# DOCX gap — corrupt -> ExtractionEngineError; valid-empty -> ""
# ---------------------------------------------------------------------------

class TestDocxCorruptHandling:

    def test_corrupt_docx_raises_engine_error_non_retryable(self):
        from app.ml import docx_extractor
        # Not a zip / not a docx package -> BadZipFile from python-docx.
        with pytest.raises(ExtractionEngineError) as exc:
            docx_extractor.extract_text_from_docx(b"this is not a docx file")
        assert exc.value.retryable is False
        assert exc.value.reason == "docx_corrupt"

    def test_valid_empty_docx_returns_empty_string(self):
        import io as _io

        import docx as _docx

        from app.ml import docx_extractor

        # Real (openable) empty .docx bytes so the zip-bomb size guard passes;
        # DocxDocument is still mocked to isolate the empty-content path.
        _buf = _io.BytesIO()
        _docx.Document().save(_buf)
        valid_empty_docx = _buf.getvalue()

        empty_doc = MagicMock()
        empty_doc.paragraphs = []
        empty_doc.tables = []
        with patch.object(docx_extractor, "DocxDocument", return_value=empty_doc):
            result = docx_extractor.extract_text_from_docx(valid_empty_docx)
        assert result == ""
